#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Combined, minimal script:
- Extracts text from PDFs/DOCX in a hard-coded folder
- Calls Azure OpenAI with your EXACT prompt
- Writes <basename>.json
- Builds the DOCX with the SAME logic as json_to_pdf.py

No CLI args; edit the CONFIG section below.
"""

import os
import re
import json
import time
import random
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import AzureOpenAI

LOGO_PATH = "./logo.png"

# Azure OpenAI config (hard-coded for now; move to env/Key Vault later)
AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT") #"https://resumeconverter.openai.azure.com/openai/deployments/gpt-4-1-mini-2025-04-14-ft-20d767268de6422099620603177f7c67-2/chat/completions?api-version=2025-01-01-preview"
AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY") 
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION") # "2024-12-01-preview"
DEPLOYMENT = os.environ.get("DEPLOYMENT") # "gpt-4-1-mini-2025-04-14-ft-20d767268de6422099620603177f7c67-2"


# =======================
# PROMPT (EXACT COPY)
# =======================
PROMPT_TEXT = r"""
"You are a resume-to-profile converter. Your task is to read the content of a resume and write a professional profile based on it."
                        "Requirements:"
                        "1. The profile must ALWAYS be written in a gender-neutral tone using they/them pronouns."
                        "2. Use plain language – no descriptive words "
                        "3. Start with past tense action verbs when explaining key responsibilities and activities"
                        "If any data is missing for any field, leave the value as an empty string ("")."
                        "4. Maintain a consistent, professional style and highlight relevant experience, achievements, and skills."
                        "3. For each work experience:"
                            "- Include the role, company, start and end dates."
                            "- Add a brief description of the company or organisation (if identifiable from context or resume content). No more than 2 sentences for each organisation/project description"
                            "- Provide key achievements and responsibilities as bullet points under Key Highlights"
                            "If any data is missing for any field, leave the value as an empty string ("")."
                            " For consultant profile with multiple projects with the same client/company: format the document in the follow:"
                            "Role "
                            "Company/Client | Project title | Mon Year – Mon Year"
                            "Project Summary in two lines max:" 
                            "Key Responsibility 1 "
                            "Key Responsibility 2 "
                            "Key Responsibility 3"

                            "Company/Client | Project title | Mon Year – Mon Year"
                            "Project Summary in two lines max: "
                            "Key Responsibility 1 "
                            "Key Responsibility 2 "
                            "Key Responsibility 3"

                            "For Consultant profile with multiple roles with the same client/company. format the document in the follow:"
                            "Client/Company"
                            "Role | Mon Year – Mon Year"
                            "Role Summary in two lines max: "
                            "Responsibility 1 "
                            "Responsibility 2" 
                            "Responsibility 3"

                            "Role | Mon Year – Mon Year"
                            "Role Summary in two lines max: "
                            "Responsibility 1" 
                            "Responsibility 2 "
                            "Responsibility 3"


                        "4. Include a detailed summary that combines strategic focus, leadership style, and industry expertise."
                            "Use this example for the summary section."
                            "This example is what we dont want- Mary is a resourceful and confident IT professional with a Bachelor of Science in Computer Engineering, with expertise in programming, system testing, problem analysis and project management both in UAT and Production. She has 15+ years of experience consulting to government and financial organisations. "
                            "Mary is a natural leader with exceptional interpersonal communications strength. She is skilled in customer service and liaising with clients and stakeholders of different nationalities and at different levels of an organisation. Mary is highly organised, with superior time management capability. She can work under pressure, with minimal supervision and has a strong willingness to learn new skills and information"
                            "Instead we want something like this: "
                            "Mary is an IT professional focusing on business analysis, business process analysis and mapping, requirements gathering, stakeholder engagement, user stories creation, facilitating and participating in Agile/Scrum ceremonies. She has expertise in problem analysis and system testing, and project management, in both UAT and Production."
                            "Mary is customer service focused.  She has proven ability liaising with clients and stakeholders. She is organised and can work under pressure, with minimal supervision."
                        "5. For Industries:"
                            "- Group them logically and, where possible, associate specific companies under each industry."
                            "Industries: Map industry → [companies]. Group logically (e.g., Telecommunications, Financial Services, Government, Healthcare, Technology)."
                            "If any data is missing for any field, leave the value as an empty string ("")."
                        "6. For Qualifications:"
                            "- Include full certificate names, issuing institutions, and dates (if available)."
                            "In Qualifications, every item must include BOTH Degree and Institution fields, even for certifications."
                            "If any data is missing for any field, leave the value as an empty string ("")."

                        "7. Keep the output structured and professional, similar to a corporate capability profile."
                        "8. If a year isn't given or is empty, use todays year"

                        "Output Rules:"     
                        "- You must output a COMPLETE, VALID JSON object with no text outside the JSON."
                        "- Do NOT cut off mid-way. Ensure all fields are closed properly and all string values are quoted and escaped."
                        "- Follow the exact JSON order and field names as the provided examples: Name, Professional Title, Industries, Qualifications, Summary, Experience, Full Work History."
                        "- Use gender-neutral language throughout (they/them)."
                        "- Ensure the tone matches an executive professional profile with descriptive company context and achievements, similar to high-quality capability documents."
                        "- Ensure out output using Australian English spelling (e.g., 'organisation', 'focussed')."
                        "- Do not use the phrase "works well without supervision"
                        "Example Output:\n"
"{\n"
"  \"Name\": \"Jane Doe\",\n"
"  \"Professional Title\": \"Data Scientist\",\n"
"  \"Industries\": {\"Technology\": [\"Google\", \"Microsoft\"]},\n"
"  \"Qualifications\": [\n"
"    {\"Degree\": \"Master of Data Science\", \"Institution\": \"University of Melbourne\"}\n"
"  ],\n"
"  \"Summary\": \"Jane is a data scientist.\",\n"
"  \"Experience\": [\n"
"    {\"Role\": \"Data Scientist\", \"Company\": \"Google\", \"Start Date\": \"Jan 2020\", "
"\"End Date\": \"Present\", \"Details\": \"Google is a global tech company.\", "
"\"Key Highlights\": [\"Led ML research\", \"Deployed NLP models\"]}\n"
"  ],\n"
"  \"Full Work History\": [\n"
"    {\"Company\": \"Google\", \"Years\": \"2020–Present\", \"Role\": \"Data Scientist\"}\n"
"  ]\n"
"}\n\n"
"Use the above structure and field names exactly. Replace the example content with "
"details extracted from the provided resume."

"""


# =======================
# DOCX builder (EXACT LOGIC)
# =======================

def add_logo_safe(paragraph, logo_path, width_cm, height_cm, align="LEFT"):
    if logo_path and os.path.exists(logo_path):
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Cm(width_cm), height=Cm(height_cm))
        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.RIGHT if align == "RIGHT" else WD_PARAGRAPH_ALIGNMENT.LEFT
        )
        return run
    return None


def set_cell_shading(cell, color="E6EEF5"):  # light grey-blue
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r_element = run._element
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)

def export_profile_to_docx(profile_data, output_path, 
                           template_path="process-resume/template.docx", 
                           logo_path=None):
    """
    Export a profile to DOCX using a template.
    
    Args:
        profile_data (dict): keys must match placeholders in the template like {{name}}, {{email}}.
        output_path (str): where to save the final DOCX.
        template_path (str): path to template.docx.
        logo_path (str): optional path to a logo image.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in profile_data.items():
            placeholder = f"{{{{{key}}}}}"  # e.g., {{name}}
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Optionally insert logo at top of document
    if logo_path:
        first_paragraph = doc.paragraphs[0]
        add_logo_safe(first_paragraph, logo_path)

    # Save the final document
    doc.save(output_path)
    return output_path


# =======================
# Helpers (extraction, model, JSON)
# =======================

def extract_text_from_pdf(filepath: Path) -> str:
    pages = []
    with pdfplumber.open(str(filepath)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_text_from_docx(filepath: Path) -> str:
    doc = Document(str(filepath))
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join((cell.text or "").strip() for cell in row.cells))
    return "\n".join(parts).strip()


def extract_text_generic(filepath: Path) -> str:
    suf = filepath.suffix.lower()
    if suf == ".pdf":
        return extract_text_from_pdf(filepath)
    if suf == ".docx":
        return extract_text_from_docx(filepath)
    raise ValueError(f"Unsupported file type: {filepath.suffix}")


def balanced_json_from_text(s: str) -> str:
    s_stripped = s.strip()
    if s_stripped.startswith("{") and s_stripped.endswith("}"):
        return s_stripped
    start = s.find("{")
    while start != -1:
        depth = 0
        for i in range(start, len(s)):
            ch = s[i]
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    candidate = s[start:i + 1]
                    try:
                        json.loads(candidate)
                        return candidate
                    except json.JSONDecodeError:
                        break
        start = s.find("{", start + 1)
    m = re.search(r"\{.*\}", s, flags=re.DOTALL)
    return m.group(0) if m else ""


# Azure OpenAI client (created once)
client = AzureOpenAI(
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
)


def call_model(resume_text: str) -> dict:
    for attempt in range(5):
        try:
            resp = client.chat.completions.create(
                model=DEPLOYMENT,
                messages=[
                    {"role": "system", "content": PROMPT_TEXT},
                    {"role": "user", "content": resume_text},
                ],
                temperature=0.2,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0,
            )
            content = (resp.choices[0].message.content or "").strip()
            json_str = balanced_json_from_text(content)
            if not json_str:
                raise ValueError("No JSON object found in model response.")
            return json.loads(json_str)
        except Exception:
            if attempt == 4:
                raise
            time.sleep((2 ** attempt) + random.uniform(0, 0.5))


def sanitize_basename(p: Path | str) -> str:
    stem = p.stem if isinstance(p, Path) else str(p)
    return re.sub(r"[^A-Za-z0-9._-]+", "_", Path(stem).stem).strip("_")

def main():
    # --- CHANGE THIS TO YOUR TEST FILE ---
    in_path = Path(r"C:\Users\arnav\PycharmProjects\resume_processor\Yohanna McLeod Resume Updated.pdf")

    if not in_path.exists():
        raise FileNotFoundError(f"Input file not found: {in_path}")
    if in_path.suffix.lower() not in {".pdf", ".docx"}:
        raise ValueError(f"Unsupported file type: {in_path.suffix} (use .pdf or .docx)")

    # Optional logo: env var wins; otherwise uses the module-level LOGO_PATH (which is None by default)
    logo_path = LOGO_PATH

    print(f"→ Extracting text from: {in_path}")
    text = extract_text_generic(in_path)
    print(f"   extracted {len(text):,} characters")

    print("→ Calling Azure OpenAI…")
    profile = call_model(text)

    base = sanitize_basename(profile.get("Name") or in_path.stem)
    out_dir = in_path.parent
    json_path = out_dir / f"{base}.json"
    docx_path = out_dir / f"{base}.docx"

    print(f"→ Writing JSON: {json_path}")
    json_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"→ Building DOCX: {docx_path}")
    export_profile_to_docx(str(json_path), str(docx_path), logo_path)


    print(f"   JSON: {json_path}")
    print(f"   DOCX: {docx_path}")


if __name__ == "__main__":
    main()
